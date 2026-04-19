"""File builder — generates PDF and DOCX resume files."""

import json
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class FileBuilder:
    """Build PDF and DOCX files from structured resume data."""

    def build_pdf(self, resume_data: dict, output_dir: str, version: int,
                  template: str = "modern", contact_info: dict = None) -> str:
        """Generate a PDF resume from structured data using WeasyPrint."""
        output_path = Path(output_dir) / f"v{version}_resume.pdf"

        try:
            html_content = self._render_html(resume_data, template, contact_info)
            from weasyprint import HTML
            HTML(string=html_content).write_pdf(str(output_path))
            logger.info(f"PDF resume saved: {output_path}")
            return str(output_path)
        except ImportError:
            logger.warning("WeasyPrint not available, creating plain text fallback")
            fallback_path = self._write_text_fallback(resume_data, str(output_path), contact_info)
            return fallback_path
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            fallback_path = self._write_text_fallback(resume_data, str(output_path), contact_info)
            return fallback_path

    def build_docx(self, resume_data: dict, output_dir: str, version: int,
                   contact_info: dict = None) -> str:
        """Generate a DOCX resume from structured data."""
        output_path = Path(output_dir) / f"v{version}_resume.docx"

        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Style defaults
            style = doc.styles["Normal"]
            font = style.font
            font.name = "Calibri"
            font.size = Pt(11)

            # Header — Name & Contact
            name = contact_info.get("name", "Candidate") if contact_info else "Candidate"
            heading = doc.add_heading(name, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if contact_info:
                contact_parts = []
                if contact_info.get("email"):
                    contact_parts.append(contact_info["email"])
                if contact_info.get("phone"):
                    contact_parts.append(contact_info["phone"])
                if contact_parts:
                    p = doc.add_paragraph(" | ".join(contact_parts))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Summary
            if resume_data.get("summary"):
                doc.add_heading("Professional Summary", level=1)
                doc.add_paragraph(resume_data["summary"])

            # Skills
            skills = resume_data.get("skills", {})
            if skills:
                doc.add_heading("Technical Skills", level=1)
                if isinstance(skills, dict):
                    for category, skill_list in skills.items():
                        if skill_list:
                            cat_name = category.replace("_", " ").title()
                            doc.add_paragraph(
                                f"{cat_name}: {', '.join(skill_list)}",
                                style="List Bullet"
                            )
                elif isinstance(skills, list):
                    doc.add_paragraph(", ".join(skills), style="List Bullet")

            # Experience
            experience = resume_data.get("experience", [])
            if experience:
                doc.add_heading("Experience", level=1)
                for exp in experience:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{exp.get('role', '')} — {exp.get('company', '')}")
                    run.bold = True
                    p.add_run(f"\n{exp.get('duration', '')}")
                    for bullet in exp.get("bullets", []):
                        doc.add_paragraph(bullet, style="List Bullet")

            # Projects
            projects = resume_data.get("projects", [])
            if projects:
                doc.add_heading("Projects", level=1)
                for proj in projects:
                    p = doc.add_paragraph()
                    run = p.add_run(proj.get("name", ""))
                    run.bold = True
                    tech = proj.get("tech_stack", "")
                    if tech:
                        p.add_run(f" | {tech}")
                    for bullet in proj.get("bullets", []):
                        doc.add_paragraph(bullet, style="List Bullet")

            # Education
            education = resume_data.get("education", [])
            if education:
                doc.add_heading("Education", level=1)
                for edu in education:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{edu.get('degree', '')} — {edu.get('institution', '')}")
                    run.bold = True
                    details = []
                    if edu.get("year"):
                        details.append(edu["year"])
                    if edu.get("cgpa"):
                        details.append(f"CGPA: {edu['cgpa']}")
                    if details:
                        p.add_run(f"\n{' | '.join(details)}")

            # Certifications
            certs = resume_data.get("certifications", [])
            if certs:
                doc.add_heading("Certifications", level=1)
                for cert in certs:
                    doc.add_paragraph(cert, style="List Bullet")

            doc.save(str(output_path))
            logger.info(f"DOCX resume saved: {output_path}")

        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            # Create a minimal text file as fallback
            output_path = Path(output_dir) / f"v{version}_resume.txt"
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(self._format_resume_text(resume_data, contact_info))

        return str(output_path)

    def _build_sections_html(self, resume_data: dict) -> str:
        """Build the HTML for skills, experience, projects, education, certifications."""
        sections = []

        # Skills section
        skills = resume_data.get("skills", {})
        if skills:
            skills_html = ""
            if isinstance(skills, dict):
                for cat, items in skills.items():
                    if items:
                        cat_name = cat.replace("_", " ").title()
                        skills_html += f'<div class="skill-cat"><strong>{cat_name}:</strong> {", ".join(items)}</div>\n'
            elif isinstance(skills, list):
                skills_html = f'<div class="skill-cat">{", ".join(skills)}</div>\n'
            if skills_html:
                sections.append(f'<h2>Technical Skills</h2>\n<div class="skills-grid">{skills_html}</div>')

        # Experience section
        experience = resume_data.get("experience", [])
        if experience:
            exp_html = ""
            for exp in experience:
                bullets = "".join(f"<li>{b}</li>" for b in exp.get("bullets", []))
                exp_html += f"""
                <div class="entry">
                    <div class="entry-header">
                        <h3>{exp.get('role', '')} <span class="company">— {exp.get('company', '')}</span></h3>
                        <span class="date">{exp.get('duration', '')}</span>
                    </div>
                    <ul>{bullets}</ul>
                </div>"""
            sections.append(f"<h2>Experience</h2>\n{exp_html}")

        # Projects section
        projects = resume_data.get("projects", [])
        if projects:
            proj_html = ""
            for proj in projects:
                bullets = "".join(f"<li>{b}</li>" for b in proj.get("bullets", []))
                tech = proj.get("tech_stack", "")
                tech_span = f' <span class="tech">| {tech}</span>' if tech else ""
                proj_html += f"""
                <div class="entry">
                    <h3>{proj.get('name', '')}{tech_span}</h3>
                    <ul>{bullets}</ul>
                </div>"""
            sections.append(f"<h2>Projects</h2>\n{proj_html}")

        # Education section
        education = resume_data.get("education", [])
        if education:
            edu_html = ""
            for edu in education:
                details = []
                if edu.get("year"):
                    details.append(edu["year"])
                if edu.get("cgpa"):
                    details.append(f"CGPA: {edu['cgpa']}")
                coursework = edu.get("relevant_coursework", [])
                coursework_html = ""
                if coursework:
                    coursework_html = f'<p class="meta">Relevant Coursework: {", ".join(coursework)}</p>'
                edu_html += f"""
                <div class="entry">
                    <div class="entry-header">
                        <h3>{edu.get('degree', '')} — {edu.get('institution', '')}</h3>
                        <span class="date">{' | '.join(details)}</span>
                    </div>
                    {coursework_html}
                </div>"""
            sections.append(f"<h2>Education</h2>\n{edu_html}")

        # Certifications section
        certs = resume_data.get("certifications", [])
        if certs:
            certs_html = "".join(f"<li>{c}</li>" for c in certs)
            sections.append(f"<h2>Certifications</h2>\n<ul>{certs_html}</ul>")

        return "\n".join(sections)

    def _render_html(self, resume_data: dict, template: str,
                     contact_info: dict = None) -> str:
        """Render resume data as HTML for PDF conversion."""
        name = contact_info.get("name", "Candidate") if contact_info else "Candidate"
        email = contact_info.get("email", "") if contact_info else ""
        phone = contact_info.get("phone", "") if contact_info else ""

        # Build the dynamic sections HTML (skills, experience, projects, education, certs)
        sections_html = self._build_sections_html(resume_data)

        # Load template file
        template_path = Path(__file__).parent.parent / "templates" / f"resume_{template}.html"
        if template_path.exists():
            with open(template_path, "r", encoding="utf-8") as f:
                html_template = f.read()

            # Replace placeholders
            html_template = html_template.replace("{{NAME}}", name)
            html_template = html_template.replace("{{EMAIL}}", email)
            html_template = html_template.replace("{{PHONE}}", phone)
            html_template = html_template.replace("{{SUMMARY}}", resume_data.get("summary", ""))

            # Inject all dynamic sections before </body>
            html_template = html_template.replace("</body>", f"{sections_html}\n</body>")

            return html_template

        # Inline fallback HTML (if no template file found)
        return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{ font-family: 'Calibri', 'Helvetica', sans-serif; margin: 40px 50px;
         font-size: 11pt; line-height: 1.4; color: #333; }}
  h1 {{ text-align: center; margin-bottom: 2px; color: #1a1a2e; font-size: 22pt; }}
  .contact {{ text-align: center; color: #666; margin-bottom: 20px; }}
  h2 {{ color: #16213e; border-bottom: 2px solid #0f3460; padding-bottom: 4px;
       font-size: 13pt; margin-top: 18px; }}
  h3 {{ margin-bottom: 2px; font-size: 11pt; }}
  .date {{ color: #888; font-style: italic; margin-top: 0; }}
  .tech {{ color: #666; font-weight: normal; font-size: 10pt; }}
  .company {{ color: #555; font-size: 10pt; }}
  .meta {{ color: #888; font-size: 9.5pt; }}
  .entry-header {{ display: flex; justify-content: space-between; align-items: baseline; }}
  .skills-grid {{ display: flex; flex-wrap: wrap; gap: 4px 20px; }}
  .skill-cat {{ margin-bottom: 4px; }}
  .skill-cat strong {{ color: #1a1a2e; }}
  ul {{ margin-top: 4px; padding-left: 20px; }}
  li {{ margin-bottom: 3px; }}
  .entry {{ margin-bottom: 12px; }}
</style></head><body>
<h1>{name}</h1>
<p class="contact">{email} | {phone}</p>
<h2>Professional Summary</h2>
<p>{resume_data.get('summary', '')}</p>
{sections_html}
</body></html>"""

    def _format_resume_text(self, resume_data: dict, contact_info: dict = None) -> str:
        """Format resume data as clean plain text."""
        lines = []
        name = contact_info.get("name", "Candidate") if contact_info else "Candidate"
        email = contact_info.get("email", "") if contact_info else ""
        phone = contact_info.get("phone", "") if contact_info else ""

        lines.append(name.upper())
        contact = " | ".join(filter(None, [email, phone]))
        if contact:
            lines.append(contact)
        lines.append("=" * 60)

        if resume_data.get("summary"):
            lines.append("\nPROFESSIONAL SUMMARY")
            lines.append("-" * 40)
            lines.append(resume_data["summary"])

        skills = resume_data.get("skills", {})
        if skills:
            lines.append("\nTECHNICAL SKILLS")
            lines.append("-" * 40)
            if isinstance(skills, dict):
                for cat, items in skills.items():
                    if items:
                        lines.append(f"  {cat.replace('_', ' ').title()}: {', '.join(items)}")
            elif isinstance(skills, list):
                lines.append(f"  {', '.join(skills)}")

        experience = resume_data.get("experience", [])
        if experience:
            lines.append("\nEXPERIENCE")
            lines.append("-" * 40)
            for exp in experience:
                lines.append(f"\n  {exp.get('role', '')} — {exp.get('company', '')}")
                lines.append(f"  {exp.get('duration', '')}")
                for bullet in exp.get("bullets", []):
                    lines.append(f"    • {bullet}")

        projects = resume_data.get("projects", [])
        if projects:
            lines.append("\nPROJECTS")
            lines.append("-" * 40)
            for proj in projects:
                tech = proj.get("tech_stack", "")
                lines.append(f"\n  {proj.get('name', '')}" + (f" | {tech}" if tech else ""))
                for bullet in proj.get("bullets", []):
                    lines.append(f"    • {bullet}")

        education = resume_data.get("education", [])
        if education:
            lines.append("\nEDUCATION")
            lines.append("-" * 40)
            for edu in education:
                lines.append(f"  {edu.get('degree', '')} — {edu.get('institution', '')}")
                details = []
                if edu.get("year"):
                    details.append(edu["year"])
                if edu.get("cgpa"):
                    details.append(f"CGPA: {edu['cgpa']}")
                if details:
                    lines.append(f"  {' | '.join(details)}")

        certs = resume_data.get("certifications", [])
        if certs:
            lines.append("\nCERTIFICATIONS")
            lines.append("-" * 40)
            for cert in certs:
                lines.append(f"  • {cert}")

        return "\n".join(lines)

    def _write_text_fallback(self, resume_data: dict, output_path: str,
                             contact_info: dict = None) -> str:
        """Write a plain text version as fallback when WeasyPrint fails.
        Returns the path to the fallback file."""
        fallback_path = output_path.replace(".pdf", ".txt")
        with open(fallback_path, "w", encoding="utf-8") as f:
            f.write(self._format_resume_text(resume_data, contact_info))
        logger.info(f"Text fallback saved: {fallback_path}")
        return fallback_path


file_builder = FileBuilder()
